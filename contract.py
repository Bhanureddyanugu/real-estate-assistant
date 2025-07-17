import streamlit as st
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from twilio.rest import Client

def send_contract_via_email(file_path, client_name, client_email):
    import yagmail
    try:
        yag = yagmail.SMTP(user="anugubhanureddy@gmail.com", password="jgfxqznxmyzbddqe")
        subject = f"Rental Agreement for Property {file_path[:6]}"
        body = f"Dear {client_name},\n\nPlease find attached your rental agreement.\n\nBest regards,\nYour Real Estate Broker"
        yag.send(to=client_email, subject=subject, contents=body, attachments=file_path)
        st.success(f"Contract sent to {client_email}")

        mailto_link = f"mailto:{client_email}?subject=" + subject.replace(" ", "%20") + "&body=" + body.replace(" ", "%20").replace("\n", "%0A")
        st.markdown(f'<a href="{mailto_link}" target="_blank">Click here to open email client</a>', unsafe_allow_html=True)
    except Exception as e:
        st.error(f"Email sending failed: {e}")

def send_whatsapp_notification(client_name, to_number):
    account_sid = "AC8ac82294d9832517499544fbc9725507"
    auth_token = "2ca56e51dff09073ca86f56055805863"
    from_whatsapp_number = "whatsapp:+14155238886"
    to_whatsapp_number = f"whatsapp:{to_number}"

    try:
        client = Client(account_sid, auth_token)
        message = client.messages.create(
            body=f"Hello {client_name}, your rental contract has been generated and sent via email. Thank you for using Dealmate.AI!",
            from_=from_whatsapp_number,
            to=to_whatsapp_number
        )
        st.success(f"WhatsApp notification sent to {to_number}")
    except Exception as e:
        st.error(f"WhatsApp sending failed: {e}")

def generate_contract_docx(client_name, property_id, city, price, email=None, whatsapp=None, contract_type="rental"):
    date_today = datetime.today().strftime("%d-%m-%Y")
    filename = f"contracts/{client_name.replace(' ', '_')}_contract_{property_id[:6]}.docx"

    if not os.path.exists("contracts"):
        os.makedirs("contracts")

    doc = Document()

    if os.path.exists("logo.png"):
        doc.add_picture("logo.png", width=Inches(1.5))

    doc.add_heading("Real Estate Agreement", 0)
    doc.add_paragraph(f"Date: {date_today}")
    doc.add_paragraph(f"Client Name: {client_name}")
    doc.add_paragraph(f"Property ID: {property_id}")
    doc.add_paragraph(f"City: {city.title()}")
    doc.add_paragraph(f"Monthly Price: ₹{price} INR")

    doc.add_heading("Terms and Conditions", level=1)
    terms = [
        f"The client agrees to {contract_type} the property listed above.",
        f"The property is located in {city.title()}.",
        f"Payment of ₹{price} is due on or before the 5th of each month.",
        "This agreement is valid for 12 months unless terminated early by mutual agreement.",
        "Property includes the listed amenities as per the official listing."
    ]
    for term in terms:
        doc.add_paragraph(term, style="List Bullet")

    doc.add_paragraph("\nSignatures:\n")
    doc.add_paragraph("___________________________        ___________________________")
    doc.add_paragraph(f"     Broker/Agent                            Client: {client_name}")
    doc.add_paragraph("\nThank you for choosing our services.")

    doc.save(filename)
    st.success(f"Contract generated: {filename}")

    if email:
        send_contract_via_email(filename, client_name, email)
    if whatsapp:
        send_whatsapp_notification(client_name, whatsapp)

def show():
    st.markdown("""
    <h2 style='color: #003B73;'>Contract Generator</h2>
    <p style='font-size: 17px; color: #374151;'>Generate rental agreements and optionally deliver them via email and WhatsApp.</p>
    <hr style='margin-bottom: 1rem;'>
    """, unsafe_allow_html=True)

    client_name = st.text_input("Client Name")
    property_id = st.text_input("Property ID")
    city = st.text_input("City")
    price = st.number_input("Monthly Price (INR)", min_value=0)
    email = st.text_input("Client Email (optional)")
    whatsapp = st.text_input("WhatsApp Number (with country code, e.g., +91...)")

    if st.button("Generate Contract"):
        if client_name and property_id and city and price:
            generate_contract_docx(client_name, property_id, city, price, email, whatsapp)
        else:
            st.warning("Please fill all required fields.")
